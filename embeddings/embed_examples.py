import os
import re
import pickle
import faiss
from sentence_transformers import SentenceTransformer
import docx
import PyPDF2
import pdfplumber

model = SentenceTransformer("all-MiniLM-L6-v2")

def read_pdf(path):
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())
    return "\n\n".join(text_parts)

def read_docx(path):
    doc = docx.Document(path)
    parts = []

    # 1. Regular paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # 2. Tables (where name/title/contact often live)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in parts:
                    parts.append(text)

    return "\n".join(parts)

def read_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text(path):
    if path.endswith(".pdf"):
        return read_pdf(path)
    elif path.endswith(".docx"):
        return read_docx(path)
    elif path.endswith(".txt"):
        return read_txt(path)
    else:
        raise ValueError(f"Unsupported file format: {path}")

def get_clean_name(filename):
    name = os.path.splitext(filename)[0]
    name = re.sub(r"(resume)?[\s_-]*(before|after)", "", name, flags=re.IGNORECASE)
    return re.sub(r"[^a-z0-9]", "", name.lower())

def extract_records(data_dir="data/before-and-after", template_file="template.docx"):
    files = os.listdir(data_dir)
    all_resumes = [f for f in files if f.lower().endswith((".pdf", ".docx", ".txt"))]
    before_files = [f for f in all_resumes if "before" in f.lower()]
    after_files = [f for f in all_resumes if "after" in f.lower()]
    after_map = {get_clean_name(f): f for f in after_files}

    # Load shared template once
    template_text = extract_text(os.path.join(data_dir, template_file)).strip()

    combined_records = []
    for bf in before_files:
        key = get_clean_name(bf)
        if key in after_map:
            before = extract_text(os.path.join(data_dir, bf)).strip()
            after = extract_text(os.path.join(data_dir, after_map[key])).strip()

            combined_text = (
    f"template:\n```text\n{template_text}\n```\n\n"
    f"before:\n```text\n{before}\n```\n\n"
    f"after:\n```text\n{after}\n```"
)
            combined_records.append(combined_text)
        else:
            print(f"❌ No after match for: {bf}")

    return combined_records

def embed_and_save(records, save_path="embeddings/faiss_index"):
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    embeddings = model.encode(records, show_progress_bar=True)

    index = faiss.IndexFlatL2(embeddings[0].shape[0])
    index.add(embeddings)

    faiss.write_index(index, save_path + ".index")
    with open(save_path + ".pkl", "wb") as f:
        pickle.dump(records, f)

def combined_chunks(examples):
    with open("data/combined_chunks.txt", "w", encoding="utf-8") as f:
        for i, text in enumerate(examples):
            f.write(f"=== Example {i+1} ===\n{text}\n\n")

if __name__ == "__main__":
    examples = extract_records()
    embed_and_save(examples)
    print(f"✅ Embedded {len(examples)} examples and saved FAISS index.")

    combined_chunks(examples)
    print("✅ Saved combined examples to data/combined_chunks.txt")
