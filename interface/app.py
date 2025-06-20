import streamlit as st
import sys
import docx
import pdfplumber
import tempfile
import os

# Ensure project root is in sys.path
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from retriever.rag_rewriter import rewrite_resume

st.set_page_config(page_title="SDR Resume Rewriter", layout="wide")
st.title("🔁 SDRPolish")

st.markdown("Upload your resume or paste raw text below. SDRPolish will rewrite it in SDR style using the SDR Career Accelerator Resume Template.")

# --- Input method ---
upload = st.file_uploader("Upload Resume (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
manual_input = st.text_area("Or paste resume text here:", height=250)

if upload:
    content = upload.read()
    filename = upload.name.lower()

    if filename.endswith(".pdf"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        with pdfplumber.open(tmp_path) as pdf:
            input_text = "\n".join(
                page.extract_text() for page in pdf.pages if page.extract_text()
            )

    elif filename.endswith(".docx"):
        doc = docx.Document(upload)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in parts:
                        parts.append(text)

        input_text = "\n".join(parts)

    elif filename.endswith(".txt"):
        input_text = upload.read().decode("utf-8")

    else:
        st.error("Unsupported file format.")
        input_text = ""
else:
    input_text = manual_input

# preview extracted text
if st.checkbox("🔍 Preview extracted resume text"):
        st.code(input_text)

# --- Submit and Rewrite ---
if input_text and st.button("🔁 Rewrite Resume"):
    with st.spinner("Rewriting with GPT..."):
        result = rewrite_resume(input_text)
    st.subheader("✅ Rewritten Resume")
    st.text_area("Result", result, height=300)
    st.download_button("⬇️ Download Rewritten Resume", result, file_name="rewritten_resume.txt")
